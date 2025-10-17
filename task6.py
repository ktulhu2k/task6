import argparse
import json
import os
import re
import sys
from abc import ABC, abstractmethod
from datetime import datetime

from docx import Document
from fpdf import FPDF, XPos, YPos
from openpyxl import Workbook


class FileSystemAnalyzer:
    """Анализирует файловую систему и возвращает структурированные данные."""

    @staticmethod
    def get_file_info(file_path):
        """
        Получает информацию о файле: имя, размер в байтах и дату последнего изменения.

        Args:
            file_path (str): Полный путь к файлу.

        Returns:
            tuple: (имя_файла, размер_в_байтах, дата_изменения_в_формате_ГГГГ-ММ-ДД ЧЧ:ММ:СС)
        """
        stat = os.stat(file_path)
        size = stat.st_size
        mtime = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
        return os.path.basename(file_path), size, mtime

    @staticmethod
    def get_folder_info(folder_path):
        """
        Получает информацию о папке: имя, фиксированную строку 'FOLDER' и дату последнего изменения.

        Args:
            folder_path (str): Полный путь к папке.

        Returns:
            tuple: (имя_папки, 'FOLDER', дата_изменения_в_формате_ГГГГ-ММ-ДД ЧЧ:ММ:СС)
        """
        stat = os.stat(folder_path)
        mtime = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
        return os.path.basename(folder_path), "FOLDER", mtime

    def analyze(self, root_path):
        """
        Рекурсивно анализирует указанную директорию и возвращает список всех файлов и папок.

        Args:
            root_path (str): Корневой путь для анализа.

        Returns:
            list[tuple]: Список элементов, где каждый элемент — кортеж вида:
                (тип_элемента, полный_путь, имя, размер_или_тип, дата_изменения)
                Типы элементов: 'folder', 'zip_folder', 'file'.
        """
        elements = []

        for dirpath, dirnames, filenames in os.walk(root_path):
            # Обрабатываем текущую папку
            name, kind, mtime = self.get_folder_info(dirpath)
            elements.append(("folder", dirpath, name, kind, mtime))

            # Обрабатываем файлы
            for filename in filenames:
                file_path = os.path.join(dirpath, filename)
                if os.path.islink(file_path):
                    continue  # пропускаем символические ссылки

                if filename.lower().endswith(".zip"):
                    name, kind, mtime = self.get_folder_info(file_path)
                    elements.append(("zip_folder", file_path, name, kind, mtime))
                else:
                    name, size, mtime = self.get_file_info(file_path)
                    elements.append(("file", file_path, name, size, mtime))

        return elements


class ReportWriter(ABC):
    """Абстрактный класс для записи отчётов в разных форматах."""

    @abstractmethod
    def write(self, elements, report_path):
        """
        Абстрактный метод для записи отчёта в файл.

        Args:
            elements (list[tuple]): Список элементов файловой структуры.
            report_path (str): Путь к выходному файлу отчёта.
        """
        pass


class CSVWriter(ReportWriter):
    def write(self, elements, report_path):
        """
        Записывает отчёт в формате CSV (Comma-Separated Values).

        Args:
            elements (list[tuple]): Список элементов файловой структуры.
            report_path (str): Путь к выходному CSV-файлу.
        """
        with open(report_path, "w", encoding="utf-8") as f:
            f.write("Тип,Имя,Размер/Тип,Дата изменения\n")
            for elem_type, _, name, size_or_type, mtime in elements:
                row_type = "Папка" if elem_type in ("folder", "zip_folder") else "Файл"
                f.write(f"{row_type},{name},{size_or_type},{mtime}\n")


class JSONWriter(ReportWriter):
    def write(self, elements, report_path):
        """
        Записывает отчёт в формате JSON (JavaScript Object Notation).

        Args:
            elements (list[tuple]): Список элементов файловой структуры.
            report_path (str): Путь к выходному JSON-файлу.
        """
        data = []
        for elem_type, _, name, size_or_type, mtime in elements:
            data.append(
                {
                    "type": "folder"
                    if elem_type in ("folder", "zip_folder")
                    else "file",
                    "name": name,
                    "size_or_type": size_or_type,
                    "modification_time": mtime,
                }
            )
        with open(report_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)


class TextWriter(ReportWriter):
    def write(self, elements, report_path):
        """
        Записывает отчёт в виде простого текстового файла (.txt или .log).

        Args:
            elements (list[tuple]): Список элементов файловой структуры.
            report_path (str): Путь к выходному текстовому файлу.
        """
        with open(report_path, "w", encoding="utf-8") as f:
            for elem_type, _, name, size_or_type, mtime in elements:
                if elem_type in ("folder", "zip_folder"):
                    f.write(
                        f"Папка: {name} | Тип: {size_or_type} | Дата изменения: {mtime}\n"
                    )
                else:
                    f.write(
                        f"Файл: {name} | Размер: {size_or_type} байт | Дата изменения: {mtime}\n"
                    )


class DOCXWriter(ReportWriter):
    def write(self, elements, report_path):
        """
        Записывает отчёт в формате DOCX (Microsoft Word).

        Args:
            elements (list[tuple]): Список элементов файловой структуры.
            report_path (str): Путь к выходному DOCX-файлу.
        """
        doc = Document()
        doc.add_heading("Отчёт о структуре файлов и папок", 0)

        for elem_type, _, name, size_or_type, mtime in elements:
            if elem_type in ("folder", "zip_folder"):
                doc.add_paragraph(
                    f"Папка: {name} | Тип: {size_or_type} | Дата изменения: {mtime}"
                )
            else:
                doc.add_paragraph(
                    f"Файл: {name} | Размер: {size_or_type} байт | Дата изменения: {mtime}"
                )

        doc.save(report_path)


class XLSXWriter(ReportWriter):
    def write(self, elements, report_path):
        """
        Записывает отчёт в формате XLSX (Microsoft Excel).

        Args:
            elements (list[tuple]): Список элементов файловой структуры.
            report_path (str): Путь к выходному XLSX-файлу.
        """
        wb = Workbook()
        ws = wb.active
        ws.title = "Отчёт о структуре файлов и папок"
        ws.append(["Тип", "Имя", "Размер/Тип", "Дата изменения"])

        for elem_type, _, name, size_or_type, mtime in elements:
            row_type = "Папка" if elem_type in ("folder", "zip_folder") else "Файл"
            ws.append([row_type, name, size_or_type, mtime])

        wb.save(report_path)


class PDFWriter(ReportWriter):
    def write(self, elements, report_path):
        """
        Записывает отчёт в формате PDF с использованием латинского шрифта Helvetica.
        Имена файлов и папок очищаются от не-ASCII символов для совместимости.

        Args:
            elements (list[tuple]): Список элементов файловой структуры.
            report_path (str): Путь к выходному PDF-файлу.
        """
        # Я долго пытался заставить pdf работать с кириллицей, честное слово
        pdf = FPDF()
        pdf.add_page()

        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(
            w=0,
            h=10,
            text="File Structure Report",
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
            align="C",
        )
        pdf.ln(6)

        pdf.set_font("Helvetica", size=9)
        page_width = 190

        for elem_type, _, name, size, mtime in elements:
            safe_name = re.sub(r"[^A-Za-z0-9._\- ()[\]{}]", "_", name)

            if elem_type in ("folder", "zip_folder"):
                line = f"[DIR]  {safe_name} | Modified: {mtime}"
            else:
                line = f"[FILE] {safe_name} | Size: {size} bytes | Modified: {mtime}"

            pdf.multi_cell(w=page_width, h=6, text=line)
            pdf.ln(1)

        pdf.output(report_path)


class ReportGenerator:
    """Основной класс для генерации отчёта."""

    def __init__(self):
        """
        Инициализирует генератор отчётов: создаёт анализатор и регистрирует все доступные форматы вывода.
        """
        self.analyzer = FileSystemAnalyzer()
        self.writers = {
            ".csv": CSVWriter(),
            ".json": JSONWriter(),
            ".txt": TextWriter(),
            ".log": TextWriter(),
            ".docx": DOCXWriter(),
            ".xlsx": XLSXWriter(),
            ".pdf": PDFWriter(),
        }

    def generate(self, input_path, output_path):
        """
        Генерирует отчёт о структуре файлов и папок по указанному пути и сохраняет его в заданном формате.

        Args:
            input_path (str): Путь к директории, которую нужно проанализировать.
            output_path (str): Путь к файлу отчёта (расширение определяет формат).

        Raises:
            FileNotFoundError: Если указанный путь не существует.
            ValueError: Если расширение файла отчёта не поддерживается.
        """
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Путь не существует: {input_path}")

        ext = os.path.splitext(output_path)[1].lower()
        writer = self.writers.get(ext)
        if not writer:
            supported = ", ".join(sorted(self.writers.keys()))
            raise ValueError(
                f"Неподдерживаемое расширение: {ext}. Поддерживаются: {supported}"
            )

        elements = self.analyzer.analyze(input_path)

        writer.write(elements, output_path)
        print(f"Отчёт сохранён: {output_path}")


def parse_arguments():
    """
    Парсит аргументы командной строки.

    Returns:
        argparse.Namespace: Объект с атрибутами path и report.
    """
    parser = argparse.ArgumentParser(
        description="Генерирует отчёт о структуре файлов и папок на жёстком диске.",
        epilog="Поддерживаемые форматы отчёта: .csv, .json, .txt, .log, .docx, .xlsx, .pdf",
    )
    parser.add_argument("--path", required=True, help="Путь к анализируемой папке")
    parser.add_argument(
        "--report", required=True, help="Путь к файлу отчёта с расширением"
    )
    return parser.parse_args()


def main():
    """
    Точка входа в программу. Парсит аргументы, создаёт генератор отчётов и запускает генерацию.
    Обрабатывает исключения и завершает программу с кодом ошибки при необходимости.
    """
    args = parse_arguments()
    generator = ReportGenerator()
    try:
        generator.generate(args.path, args.report)
    except Exception as e:
        print(f"Ошибка: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
